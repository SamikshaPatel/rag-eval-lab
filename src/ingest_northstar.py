"""
NORTHSTAR INGESTION
===================
Loads data/Northstar_Digital_Bank.docx, chunks it, embeds with the same
nomic-embed-text model, and persists to a separate Chroma DB at
./chroma_db_northstar/ so the Northstar and Zephyr corpora stay isolated.

Run once before eval_langsmith_northstar.py:
    python src/ingest_northstar.py
"""

from docx import Document
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma

from config import (
    EMBED_MODEL, CHUNK_SIZE, CHUNK_OVERLAP,
    NORTHSTAR_DB_PATH, NORTHSTAR_DATA_PATH,
)


def load_docx(path: str) -> str:
    """Extract all paragraph text from a .docx file as a single string."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def main():
    print(f"Loading {NORTHSTAR_DATA_PATH} ...")
    text = load_docx(NORTHSTAR_DATA_PATH)
    print(f"  {len(text):,} characters loaded.")

    # Wrap in a LangChain Document so the splitter can process it
    raw_doc = LCDocument(page_content=text, metadata={"source": NORTHSTAR_DATA_PATH})

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents([raw_doc])
    print(f"  Split into {len(chunks)} chunks (size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP}).")

    print(f"Embedding with {EMBED_MODEL} and persisting to {NORTHSTAR_DB_PATH} ...")
    embeddings = OllamaEmbeddings(model=EMBED_MODEL)
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=NORTHSTAR_DB_PATH,
    )
    print(f"Ingestion complete — {len(chunks)} chunks stored in {NORTHSTAR_DB_PATH}.")
    print("You can now run:  python src/eval_langsmith_northstar.py")


if __name__ == "__main__":
    main()
